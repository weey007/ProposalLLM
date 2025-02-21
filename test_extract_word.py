import unittest
from pathlib import Path
from docx import Document
from docx.table import _Cell
from Extract_Word import DocumentProcessor

class TestDocumentProcessor(unittest.TestCase):
    def setUp(self):
        """测试前的设置"""
        print("\n开始测试环境设置...")
        self.processor = DocumentProcessor()
        self.test_dir = Path(__file__).parent / "test_files"
        self.test_dir.mkdir(exist_ok=True)
        
        # 创建测试用Word文档
        self.test_doc = Document()
        self.test_doc.add_heading("测试标题1", level=1)
        self.test_doc.add_paragraph("测试内容1")
        self.test_doc.add_heading("测试标题2", level=2)
        self.test_doc.add_paragraph("测试内容2")
        self.test_doc_path = self.test_dir / "test.docx"
        self.test_doc.save(str(self.test_doc_path))
        print("测试环境设置完成")

    def tearDown(self):
        """测试后的清理"""
        print("\n开始清理测试环境...")
        # 删除测试文件
        if self.test_doc_path.exists():
            self.test_doc_path.unlink()
        if self.test_dir.exists():
            self.test_dir.rmdir()
        print("测试环境清理完成")

    def test_get_file_name(self):
        """测试文件名生成功能"""
        print("\n开始测试文件名生成功能...")
        # 测试基本文件名生成
        self.processor.version = [1, 0, 0]
        result = self.processor.get_file_name("测试标题")
        print(f"基本文件名生成测试: {result}")
        self.assertEqual(result, "1- 测试标题")

        # 测试特殊字符处理
        self.processor.version = [1, 1, 0]
        result = self.processor.get_file_name("测试*标题#")
        print(f"特殊字符处理测试: {result}")
        self.assertEqual(result, "1.1- 测试标题")
        print("文件名生成功能测试完成")

    def test_update_version(self):
        """测试版本号更新功能"""
        print("\n开始测试版本号更新功能...")
        # 测试一级标题
        self.processor.version = [0, 0, 0]
        self.processor.update_version("1")
        print(f"一级标题版本号更新: {self.processor.version}")
        self.assertEqual(self.processor.version, [1, 0, 0])

        # 测试二级标题
        self.processor.update_version("2")
        print(f"二级标题版本号更新: {self.processor.version}")
        self.assertEqual(self.processor.version, [1, 1, 0])

        # 测试三级标题
        self.processor.update_version("3")
        print(f"三级标题版本号更新: {self.processor.version}")
        self.assertEqual(self.processor.version, [1, 1, 1])
        print("版本号更新功能测试完成")

    def test_set_cell_font(self):
        """测试单元格字体设置功能"""
        print("\n开始测试单元格字体设置功能...")
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell: _Cell = table.cell(0, 0)
        cell.text = "测试文本"
        
        self.processor.set_cell_font(cell, "宋体")
        print("已设置单元格字体为宋体")
        
        # 验证字体设置
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                self.assertEqual(run.font.name, "宋体")
        print("单元格字体设置功能测试完成")

if __name__ == '__main__':
    unittest.main()
