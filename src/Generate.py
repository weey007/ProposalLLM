import io
import json
import os
import re
from pathlib import Path
from typing import Optional, List, Dict, Any

import docx
import openai
import openpyxl
import requests
from docx.shared import Cm
from dotenv import load_dotenv

# 定义项目根目录和其他目录
ROOT_DIR = Path(__file__).parent
DATA_DIR = ROOT_DIR / "data"
INPUT_DIR = DATA_DIR / "input"
OUTPUT_DIR = DATA_DIR / "output"
TEMPLATES_DIR = DATA_DIR / "templates"
EXAMPLES_DIR = ROOT_DIR / "examples"

# 创建必要的目录
for dir_path in [DATA_DIR, INPUT_DIR, OUTPUT_DIR, TEMPLATES_DIR, EXAMPLES_DIR]:
    dir_path.mkdir(parents=True, exist_ok=True)

# 加载环境变量
load_dotenv()


class Config:
    """配置类，用于管理所有配置项"""

    # 文件路径配置
    EXCEL_FILE = INPUT_DIR / "需求对应表.xlsx"
    WORD_FILE = INPUT_DIR / "标书内容.docx"
    TEMPLATE_FILE = TEMPLATES_DIR / "Template.docx"

    # API配置
    BAIDU_API_KEY = os.getenv('BAIDU_API_KEY', '')
    BAIDU_SECRET_KEY = os.getenv('BAIDU_SECRET_KEY', '')
    OPENAI_API_KEY = os.getenv('OPENAI_API_KEY', '')
    OPENAI_API_BASE = os.getenv('OPENAI_API_BASE', 'https://api.openai.com/v1')

    # API选择
    USE_BAIDU = os.getenv('USE_BAIDU', 'true').lower() == 'true'

    # OpenAI配置
    OPENAI_MODEL = os.getenv('OPENAI_MODEL', 'gpt-4')
    OPENAI_TEMPERATURE = float(os.getenv('OPENAI_TEMPERATURE', '0.7'))
    OPENAI_MAX_TOKENS = int(os.getenv('OPENAI_MAX_TOKENS', '1500'))

    # Word文档配置
    MAX_WIDTH_CM = 14.0

    # 提示词配置
    PROMPT_ANSWER = """现在有个问答，比选文件要求和比选申请人应答，我给你举个例子 ，比如比选文件要求：支持可视化创建不同类型数据源，
    包括但不限于：传统数据库、文件系统、消息队列、SaaS API，NoSQL等、必选申请人回答的是：完全支持。系统支持数据源配置化管理，
    数据源、数据目标的信息可界面化管理。支持新增、修改、删除等配置管理功能，支持搜索功能。你学习一下我的风格。
    现在我是比选申请人，请严格按照我的风格来回答，请注意我回答的格式：首先是'完全支持'，然后说'系统支持什么什么', 
    这个过程需要你按照问题回答，不要跑题。例如，输入我的整体回答就变成了：'完全支持。系统支持数据源配置化管理，
    数据源、数据目标的信息可界面化管理。支持新增、修改、删除等配置高级管理功能，全面支持搜索功能。'以下是输入文字："""

    PROMPT_CONTENT = """你是一个大数据平台的专业产品售前，请针对这一需求给出800字的产品功能介绍，
    不要开头和总结，直接写产品功能，不需要用markdown格式，直接文本格式+特殊项目符号输出即可，需求如下:"""

    PROMPT_TITLE = """你是一个专业作者，请把以下这段文字变为10字以内不带细节内容和标点和解释的文字，
    直接给出结果不要'简化为'这种返回："""

    # 标题配置
    MORE_SECTION = 1
    RE_GENERATE_TEXT = 0
    DDD_ANSWER = 1
    KEY_FLAG = 0
    LEVEL1 = 'heading 1'
    LEVEL2 = 'heading 2'
    LAST_HEADING_1 = 2
    LAST_HEADING_2 = 0
    LAST_HEADING_3 = 0


class BaiduAPI:
    """百度API调用类"""

    @staticmethod
    def get_access_token() -> str:
        """获取百度API访问令牌"""
        url = "https://aip.baidubce.com/oauth/2.0/token"
        params = {
            "grant_type": "client_credentials",
            "client_id": Config.BAIDU_API_KEY,
            "client_secret": Config.BAIDU_SECRET_KEY
        }

        try:
            resp = requests.post(url, params=params).json()
            if 'error' in resp:
                raise Exception(f"获取token失败: {resp.get('error_description', '未知错误')}")
            return str(resp.get("access_token"))
        except Exception as e:
            print(f"获取百度API token失败: {str(e)}")
            raise

    @classmethod
    def call_api(cls, prompt: str) -> str:
        """调用百度API"""
        url = f"https://aip.baidubce.com/rpc/2.0/ai_custom/v1/wenxinworkshop/chat/ernie_speed?access_token={cls.get_access_token()}"

        payload = {
            "messages": [{"role": "user", "content": prompt}]
        }

        try:
            response = requests.post(url, json=payload)
            data = response.json()
            if 'error_code' in data:
                raise Exception(f"API调用失败: {data.get('error_msg', '未知错误')}")
            return data['result']
        except Exception as e:
            print(f"调用百度API失败: {str(e)}")
            raise


class OpenAIAPI:
    """OpenAI API调用类"""

    @staticmethod
    def initialize():
        """初始化OpenAI配置"""
        if not Config.OPENAI_API_KEY:
            raise ValueError("OpenAI API密钥未配置")
        openai.api_key = Config.OPENAI_API_KEY
        openai.api_base = Config.OPENAI_API_BASE

    @classmethod
    def call_api(cls, prompt: str) -> str:
        """调用OpenAI API"""
        try:
            cls.initialize()
            response = openai.ChatCompletion.create(
                model=Config.OPENAI_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=Config.OPENAI_TEMPERATURE,
                max_tokens=Config.OPENAI_MAX_TOKENS
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            print(f"调用OpenAI API失败: {str(e)}")
            raise


class AIService:
    """AI服务类，处理所有AI相关的操作"""

    @staticmethod
    def get_ai_provider():
        """获取AI提供商"""
        return BaiduAPI if Config.USE_BAIDU else OpenAIAPI

    @classmethod
    def generate_solution(cls, content: str) -> str:
        """生成解决方案"""
        ai_provider = cls.get_ai_provider()
        prompt = f"{Config.PROMPT_CONTENT} {content}"
        return ai_provider.call_api(prompt)

    @classmethod
    def shorten_text(cls, text: str) -> str:
        """将文本缩减为标题"""
        ai_provider = cls.get_ai_provider()
        prompt = f"{Config.PROMPT_TITLE}'{text}'"
        result = ai_provider.call_api(prompt)
        cleaned_title = result.replace("。", "")

        if Config.KEY_FLAG == 1:
            if '★' in text and '★' not in cleaned_title:
                cleaned_title = f"★{cleaned_title}"
            elif '▲' in text and '▲' not in cleaned_title:
                cleaned_title = f"▲{cleaned_title}"
        return cleaned_title

    @classmethod
    def optimize_description(cls, text: str) -> str:
        """优化需求说明"""
        ai_provider = cls.get_ai_provider()
        prompt = f"{Config.PROMPT_ANSWER}'{text}'"
        return ai_provider.call_api(prompt)


class DocumentProcessor:
    """文档处理类"""

    @staticmethod
    def get_image_from_run(run) -> Optional[tuple]:
        """从run中提取图片"""
        try:
            for item in run._r.drawing_lst:
                if item.graphic.graphicData.pic is not None:
                    rId = item.graphic.graphicData.pic.blipFill.blip.embed
                    image = run.part.get_or_add_image(rId)
                    return image.blob, image.content_type
        except Exception as e:
            print(f"提取图片失败: {str(e)}")
        return None

    @staticmethod
    def iter_block_items(parent):
        """遍历文档块"""
        if isinstance(parent, docx.document.Document):
            parent_elm = parent.element.body
        else:
            parent_elm = parent._element

        for child in parent_elm.iterchildren():
            if isinstance(child, docx.oxml.text.paragraph.CT_P):
                yield docx.text.paragraph.Paragraph(child, parent)
            elif isinstance(child, docx.oxml.table.CT_Tbl):
                yield docx.table.Table(child, parent)


class ExcelProcessor:
    """Excel处理类"""

    @staticmethod
    def load_excel(file_path: str) -> openpyxl.Workbook:
        """加载Excel文件"""
        try:
            return openpyxl.load_workbook(file_path)
        except Exception as e:
            print(f"加载Excel文件失败: {str(e)}")
            raise

    @staticmethod
    def get_sheet(workbook: openpyxl.Workbook) -> openpyxl.Worksheet:
        """获取Excel工作表"""
        try:
            return workbook.active
        except Exception as e:
            print(f"获取Excel工作表失败: {str(e)}")
            raise


class WordProcessor:
    """Word文档处理类"""

    @staticmethod
    def load_word(file_path: str) -> docx.Document:
        """加载Word文件"""
        try:
            return docx.Document(file_path)
        except Exception as e:
            print(f"加载Word文件失败: {str(e)}")
            raise

    @staticmethod
    def save_word(document: docx.Document, file_path: str) -> None:
        """保存Word文件"""
        try:
            document.save(file_path)
        except Exception as e:
            print(f"保存Word文件失败: {str(e)}")
            raise

    @staticmethod
    def find_word_file(g_column_value: str) -> str:
        """根据G列值找到对应的Word文件"""
        try:
            # 移除可能存在的文件扩展名
            g_column_value = str(g_column_value).split('.')[0]

            # 在当前目录查找匹配的文件
            for file in os.listdir():
                if file.endswith('.docx'):
                    file_prefix = file.split('-')[0]
                    if file_prefix == g_column_value:
                        return file
            raise FileNotFoundError(f"未找到对应的Word文件: {g_column_value}")
        except Exception as e:
            print(f"查找Word文件失败: {str(e)}")
            raise

    @staticmethod
    def copy_content_with_images(doc: docx.Document, target_doc: docx.Document) -> None:
        """复制文本和图片到目标文档"""
        try:
            for element in DocumentProcessor.iter_block_items(doc):
                if isinstance(element, docx.text.paragraph.Paragraph):
                    p = target_doc.add_paragraph()
                    for run in element.runs:
                        new_run = p.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic

                        # 处理图片
                        image_data = DocumentProcessor.get_image_from_run(run)
                        if image_data:
                            blob, content_type = image_data
                            if blob:
                                image_stream = io.BytesIO(blob)
                                width = None

                                # 检查图片宽度是否超过最大宽度
                                try:
                                    from PIL import Image
                                    with Image.open(image_stream) as img:
                                        width_cm = img.width / 37.795275591  # 将像素转换为厘米
                                        if width_cm > Config.MAX_WIDTH_CM:
                                            width = Cm(Config.MAX_WIDTH_CM)
                                    image_stream.seek(0)
                                except Exception as e:
                                    print(f"处理图片尺寸时出错: {str(e)}")

                                if width:
                                    p.add_run().add_picture(image_stream, width=width)
                                else:
                                    p.add_run().add_picture(image_stream)
                elif isinstance(element, docx.table.Table):
                    target_doc.add_table(element._tbl)
        except Exception as e:
            print(f"复制内容时出错: {str(e)}")
            raise


def main():
    """主函数"""
    try:
        # 检查环境变量是否已配置
        if not Config.BAIDU_API_KEY or not Config.BAIDU_SECRET_KEY:
            raise ValueError("请在.env文件中配置BAIDU_API_KEY和BAIDU_SECRET_KEY")

        # 加载配置
        config = Config()

        # 加载Excel文件
        excel_file = Config.EXCEL_FILE
        if not os.path.exists(excel_file):
            raise FileNotFoundError(f"未找到Excel文件: {excel_file}")

        excel_processor = ExcelProcessor()
        workbook = excel_processor.load_excel(excel_file)
        sheet = excel_processor.get_sheet(workbook)

        # 加载Word文件
        word_file = Config.WORD_FILE
        word_processor = WordProcessor()
        document = word_processor.load_word(word_file)

        # 处理Excel数据
        for row in sheet.iter_rows(min_row=2):  # 从第二行开始，跳过表头
            b_column_content = row[1].value
            c_column_content = row[2].value
            g_column_value = row[6].value

            # 调用百度API，缩减C列内容
            shortened_title = AIService.shorten_text(c_column_content)

            # 调用OpenAI API，对C列做应答
            optimized_description = AIService.generate_solution(c_column_content)

            # 将优化后的说明写入E列
            row[4].value = optimized_description
            # 将需求写入到D列
            row[3].value = shortened_title

            # 根据MoreSection逻辑生成标题，并处理标题二和标题三的级联关系
            if config.MORE_SECTION == 1:
                if b_column_content:
                    config.LAST_HEADING_2 += 1
                    config.LAST_HEADING_3 = 1
                    document.add_heading(f" {b_column_content}", level=2)

                    # 将当前标题二章节号写入F列
                    row[5].value = f"{config.LAST_HEADING_1}.{config.LAST_HEADING_2}.{config.LAST_HEADING_3}"

                    # 生成标题三的小节，标题基于C列内容
                    document.add_heading(f"{shortened_title}", level=3)
                else:
                    config.LAST_HEADING_3 += 1
                    document.add_heading(f" {shortened_title}", level=3)

                    # 将当前标题三章节号写入F列
                    row[5].value = f"{config.LAST_HEADING_1}.{config.LAST_HEADING_2}.{config.LAST_HEADING_3}"
            else:
                config.LAST_HEADING_2 += 1
                document.add_heading(f"{config.LAST_HEADING_2}. {shortened_title}", level=2)

                # 将当前标题二章节号写入F列
                row[5].value = f"{config.LAST_HEADING_1}.{config.LAST_HEADING_2}"

            # 将G列内容转换为整数，作为要打开的Word文档名
            try:
                x_word_file = f"{g_column_value}.docx"
                # 检查对应的Word文件是否存在
                if os.path.exists(x_word_file):
                    x_document = word_processor.load_word(x_word_file)
                    # 复制文件和图片
                    for block in DocumentProcessor.iter_block_items(x_document):
                        if isinstance(block, docx.text.paragraph.Paragraph):
                            paragraph = block
                            if paragraph.style.name.startswith('Heading'):
                                if paragraph.style.name == 'Heading 2':
                                    config.LAST_HEADING_2 += 1
                                    config.LAST_HEADING_3 = 1
                                elif paragraph.style.name == 'Heading 3':
                                    config.LAST_HEADING_3 += 1
                                document.add_heading(paragraph.text, level=int(paragraph.style.name[-1]))
                            else:
                                document.add_paragraph(paragraph.text, style='Normal')
                        elif isinstance(block, docx.table.Table):
                            table = block
                            document.add_table(rows=len(table.rows), cols=len(table.columns))
                            for i, row in enumerate(table.rows):
                                for j, cell in enumerate(row.cells):
                                    document.tables[-1].cell(i, j).text = cell.text
            except ValueError:
                print(f"Error converting G column to integer: {g_column_value}")
                continue

        # 保存更新后的Excel文件
        workbook.save(excel_file)

        # 保存最终Word文档
        word_processor.save_word(document, word_file)

        print("Process completed!")

    except Exception as e:
        print(f"程序执行出错: {str(e)}")
        raise


if __name__ == "__main__":
    main()
