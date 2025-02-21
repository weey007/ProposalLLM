"""
Word文档处理工具

这个模块提供了将Word文档按照标题层级拆分成多个子文档的功能。
支持处理文本、表格、图片等内容，并保持原有格式。

主要功能：
- 按照标题层级拆分文档
- 保持原始格式（包括中文字体）
- 处理表格和图片
- 自动调整图片大小
"""

from typing import List, Tuple, Optional, Union, BinaryIO
from docx import Document
from docx.document import Document as _Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.shared import Cm
import io
import os

class DocumentProcessor:
    """处理Word文档的主类"""
    
    MAX_WIDTH_CM = 14.0  # 最大宽度（厘米）

    def __init__(self):
        self.version = [0, 0, 0]

    def get_file_name(self, heading_text: str) -> str:
        """生成文件名
        
        Args:
            heading_text: 标题文本
            
        Returns:
            格式化后的文件名
        """
        version = self.version.copy()
        while version[-1] == 0:
            version.pop()
        
        # 清理标题文本，只保留字母数字和特定符号
        clean_text = ''.join(c for c in heading_text if c.isalnum() or c in (' ', '-', '_')).strip()
        return f"{'.'.join(map(str, version))}- {clean_text}"

    def update_version(self, level: str) -> None:
        """更新版本号
        
        Args:
            level: 标题级别（'1', '2', '3'）
        """
        while len(self.version) < 3:
            self.version.append(0)
        
        if level == '1':
            self.version[0] += 1
            self.version[1:] = [0, 0]
        elif level == '2':
            self.version[1] += 1
            self.version[2] = 0
        elif level == '3':
            self.version[2] += 1

    @staticmethod
    def set_cell_font(cell: _Cell, font_name: str = '宋体') -> None:
        """设置单元格字体
        
        Args:
            cell: 表格单元格
            font_name: 字体名称
        """
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def process_table(self, table: Table) -> Tuple[str, List[List[str]]]:
        """处理表格
        
        Args:
            table: Word表格对象
            
        Returns:
            包含表格数据的元组
        """
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = ''
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if not self.get_image_from_run(run):
                            cell_text += run.text.strip()
                row_data.append(cell_text)
            table_data.append(row_data)
        return ('table', table_data)

    @staticmethod
    def get_image_from_run(run) -> Optional[Tuple[BinaryIO, str, float, float]]:
        """从运行对象中提取图片
        
        Args:
            run: Word运行对象
            
        Returns:
            图片数据元组或None
        """
        drawing_elements = run._element.xpath('.//a:blip')
        if not drawing_elements:
            return None

        embed = drawing_elements[0].get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        image_part = run.part.related_parts[embed]
        image_data = image_part.blob
        
        width_cm = height_cm = None
        drawing_element = run._element.xpath('.//wp:extent')
        if drawing_element:
            cx = int(drawing_element[0].get('cx'))
            cy = int(drawing_element[0].get('cy'))
            width_cm = cx / 914400 * 2.54
            height_cm = cy / 914400 * 2.54

        image_stream = io.BytesIO(image_data)
        image_stream.name = os.path.basename(image_part.partname)
        return image_stream, image_stream.name, width_cm, height_cm

    def save_content_to_new_doc(self, content: List[Union[str, Tuple]], heading_text: str) -> None:
        """保存内容到新的Word文档
        
        Args:
            content: 要保存的内容列表
            heading_text: 标题文本
        """
        file_name = f"{self.get_file_name(heading_text)}.docx"
        doc = Document()
        
        for item in content:
            if isinstance(item, str):
                self._add_text_paragraph(doc, item)
            elif isinstance(item, tuple):
                self._process_content_item(doc, item)

        doc.save(file_name)

    def _add_text_paragraph(self, doc: _Document, text: str) -> None:
        """添加文本段落
        
        Args:
            doc: Word文档对象
            text: 文本内容
        """
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _process_content_item(self, doc: _Document, item: Tuple) -> None:
        """处理内容项
        
        Args:
            doc: Word文档对象
            item: 内容项元组
        """
        item_type = item[0]
        if item_type == 'table':
            self._add_table(doc, item[1])
        elif item_type == 'image':
            self._add_image(doc, *item[1:])
        elif item_type == 'list':
            self._add_list_item(doc, item[1])

    def _add_table(self, doc: _Document, table_data: List[List[str]]) -> None:
        """添加表格
        
        Args:
            doc: Word文档对象
            table_data: 表格数据
        """
        if not table_data:
            return
            
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        for i, row in enumerate(table_data):
            for j, cell_text in enumerate(row):
                cell = table.cell(i, j)
                cell.text = cell_text
                self.set_cell_font(cell)

        # 添加表格边框
        tbl = table._element
        tbl_pr = tbl.tblPr
        tbl_borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

    def _add_image(self, doc: _Document, image_stream: BinaryIO, 
                  image_name: str, width_cm: float, height_cm: float) -> None:
        """添加图片
        
        Args:
            doc: Word文档对象
            image_stream: 图片数据流
            image_name: 图片名称
            width_cm: 宽度（厘米）
            height_cm: 高度（厘米）
        """
        if width_cm > self.MAX_WIDTH_CM:
            scale_factor = self.MAX_WIDTH_CM / width_cm
            width_cm = self.MAX_WIDTH_CM
            height_cm = height_cm * scale_factor
        doc.add_paragraph().add_run().add_picture(
            image_stream, width=Cm(width_cm), height=Cm(height_cm))

    def _add_list_item(self, doc: _Document, text: str) -> None:
        """添加列表项
        
        Args:
            doc: Word文档对象
            text: 列表项文本
        """
        paragraph = doc.add_paragraph(text, style='List Paragraph')
        for run in paragraph.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    @staticmethod
    def iter_block_items(parent: Union[_Document, Paragraph, Table]):
        """遍历文档块
        
        Args:
            parent: 父对象
            
        Yields:
            段落或表格对象
        """
        parent_elm = parent.element.body if isinstance(parent, _Document) else parent._element
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def process_document(self, docx_path: str) -> None:
        """处理Word文档
        
        Args:
            docx_path: Word文档路径
        """
        doc = Document(docx_path)
        content_between_headings = []
        current_heading_text = None
        
        for block in self.iter_block_items(doc):
            if isinstance(block, Paragraph):
                self._process_paragraph(block, content_between_headings)
            elif isinstance(block, Table):
                content_between_headings.append(self.process_table(block))

        if content_between_headings and current_heading_text:
            self.save_content_to_new_doc(content_between_headings, current_heading_text)

    def _process_paragraph(self, para: Paragraph, content: List) -> None:
        """处理段落
        
        Args:
            para: 段落对象
            content: 内容列表
        """
        if para.style.name.startswith('Heading'):
            if content:
                self.save_content_to_new_doc(content, para.text)
                content.clear()
            
            heading_level = para.style.name.split(' ')[-1]
            self.update_version(heading_level)
            return

        paragraph_text = ""
        for run in para.runs:
            image_data = self.get_image_from_run(run)
            if image_data:
                content.append(('image',) + image_data)
            else:
                text = run.text.strip()
                if text:
                    paragraph_text += text

        if para.style.name == 'List Paragraph' or para._element.xpath('.//w:numPr'):
            content.append(('list', paragraph_text))
        elif paragraph_text:
            content.append(paragraph_text)

def main():
    """主函数"""
    processor = DocumentProcessor()
    docx_file = 'Template.docx'  # 替换为实际的文档路径
    processor.process_document(docx_file)

if __name__ == '__main__':
    main()
